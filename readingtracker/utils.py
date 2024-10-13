import os
from django.conf import settings
from django.core.exceptions import ValidationError
import PyPDF2
from docx import Document

def validate_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            # Check if it has at least one page
            if len(pdf_reader.pages) < 1:
                raise ValidationError("The file is not a valid PDF.")
        return True
    except Exception:
        raise ValidationError("The file is not a valid PDF.")

def validate_docx(file_path):
    try:
        doc = Document(file_path)
        # Check if it has at least one paragraph
        if len(doc.paragraphs) < 1:
            raise ValidationError("The file is not a valid DOCX.")
        return True
    except Exception:
        raise ValidationError("The file is not a valid DOCX.")

def validate_file_type(file):
    file_path = os.path.join(settings.MEDIA_ROOT, file.name)

    # Save the file temporarily for validation
    with open(file_path, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)

    try:
        # Check the file type
        if file.name.endswith('.pdf'):
            is_valid = validate_pdf(file_path)
        elif file.name.endswith('.docx'):
            is_valid = validate_docx(file_path)
        else:
            raise ValidationError("Unsupported file type. Please upload a PDF or DOCX file.")

        return is_valid
    finally:
        # Clean up the temporary file
        if os.path.exists(file_path):
            os.remove(file_path)

def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + '\n'
    return text

def extract_text_from_docx(file_path):
    text = ''
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text
